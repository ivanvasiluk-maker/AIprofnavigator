from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from html import escape, unescape
from pathlib import Path
import re
import tempfile
import textwrap
from urllib.parse import unquote, urlparse

from config import settings


@dataclass
class ReportMeta:
    user_name: str
    country: str
    mode: str
    created_at: str
    profile_version: str = ""


def _resolve_unicode_font_path() -> Path | None:
    candidates: list[Path] = []
    if settings.report_pdf_font_path:
        candidates.append(Path(settings.report_pdf_font_path))

    candidates.extend([
        Path("fonts/DejaVuSans.ttf"),
        Path("fonts/NotoSans-Regular.ttf"),
        Path("fonts/ArialUnicodeMS.ttf"),
        Path("fonts/SegoeUI.ttf"),
        Path("C:/Windows/Fonts/DejaVuSans.ttf"),
        Path("C:/Windows/Fonts/NotoSans-Regular.ttf"),
        Path("C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arialuni.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
    ])
    for path in candidates:
        if path.exists() and path.is_file():
            return path.resolve()
    return None


def _safe_text(value: object, default: str = "-") -> str:
    if value is True:
        return "есть"
    if value is False:
        return "нет"
    text = str(value or "").strip()
    if text.casefold() in {"unknown", "none", "null", "undefined"}:
        return "не указано"
    return text if text else default


def _list_items(items: object, fallback: str = "-") -> list[str]:
    if not isinstance(items, list):
        return [fallback]
    cleaned = [str(item).strip() for item in items if str(item).strip()]
    return cleaned or [fallback]


def _level_label(value: object) -> str:
    normalized = str(value or "").strip().lower()
    labels = {
        "high": "высокий",
        "medium": "средний",
        "low": "низкий",
    }
    return labels.get(normalized, _safe_text(value, "не уточнено"))


def _resource_human_message(value: object) -> str:
    normalized = str(value or "").strip().lower()
    if normalized == "low":
        return (
            "Сейчас ресурс ограничен.\n"
            "Вам важнее не брать на себя большой карьерный разворот, а выбрать путь, который дает доход "
            "и не требует резко учиться всему с нуля."
        )
    if normalized == "medium":
        return (
            "Сейчас ресурс частично ограничен.\n"
            "Лучше двигаться короткими шагами: сначала стабильный доход и понятный ритм, затем расширять траекторию."
        )
    if normalized == "high":
        return (
            "Сейчас ресурс устойчивый.\n"
            "Вы можете брать более амбициозные шаги, но сохранять контроль: измеримые действия и проверка гипотез на рынке."
        )
    return (
        "Ресурс пока нельзя оценить точно.\n"
        "Нужно уточнить, насколько стабильно у вас хватает времени, сил и темпа на регулярные карьерные действия."
    )


def _integration_human_message(value: object) -> str:
    normalized = str(value or "").strip().lower()
    if normalized == "low":
        return (
            "Интеграция пока начальная.\n"
            "Нужен щадящий маршрут: меньше зависимости от широких контактов и сложных локальных процессов на старте."
        )
    if normalized == "medium":
        return (
            "Интеграция пока частичная.\n"
            "Вы уже живете и работаете в стране, но язык, местные контакты и понимание рынка еще требуют укрепления."
        )
    if normalized == "high":
        return (
            "Интеграция уже устойчивая.\n"
            "Можно делать ставку на роли с более высокой планкой входа и активнее использовать локальный нетворк."
        )
    return (
        "Интеграцию пока нельзя оценить точно.\n"
        "Мы знаем, что вы уже живете в стране, но пока не понимаем, насколько уверенно вы используете язык, "
        "контакты и местные сервисы."
    )


def _professional_core_summary(report: dict) -> str:
    digital_human = report.get("digital_human") if isinstance(report.get("digital_human"), dict) else {}
    not_reset = report.get("what_not_reset") if isinstance(report.get("what_not_reset"), list) else []
    snapshot = " ".join(str(item) for item in not_reset[:6]).lower().replace("ё", "е")
    state_blob = " ".join(
        [
            str(digital_human.get("current_state", "")),
            str(digital_human.get("main_asset", "")),
            snapshot,
        ]
    ).lower().replace("ё", "е")

    worker_markers = ["плитк", "гипсокарт", "ремонт", "строит", "мебел", "рукам", "отделк"]
    if any(marker in state_blob for marker in worker_markers):
        return (
            "Вы не начинаете с нуля. Ваш основной капитал — практический опыт, способность работать руками, "
            "доводить задачу до результата и общаться с заказчиком. "
            "Сейчас вам нужен не абстрактный «новый старт», а перевод уже имеющегося опыта в понятный доход в новой стране."
        )

    return (
        "Вы не начинаете с нуля. У вас уже есть профессиональный капитал: практический опыт, рабочая дисциплина, "
        "умение доводить задачи до результата и взаимодействовать с людьми. "
        "Сейчас нужен не «новый старт с пустого места», а точная упаковка и перевод вашего опыта в понятный доход в новой стране."
    )


def _detect_country(report: dict) -> str:
    decision = report.get("career_decision") if isinstance(report.get("career_decision"), dict) else {}
    snapshot = report.get("profile_snapshot") if isinstance(report.get("profile_snapshot"), dict) else {}
    country_name = (
        str(decision.get("country_name") or snapshot.get("country_name") or report.get("country_name") or "").strip()
        or str(snapshot.get("country_code") or "").strip()
        or str((snapshot.get("route_context") or {}).get("country") or "").strip()
    )
    if country_name:
        return country_name
    profile = str((report.get("digital_human") or {}).get("current_state", "")).lower()
    if "польш" in profile or "poland" in profile:
        return "Poland"
    if "беларус" in profile:
        return "Belarus"
    return "Не уточнено"


def _detect_mode(report: dict) -> str:
    """Survival requires explicit user-confirmed financial urgency evidence."""
    dh = report.get("digital_human") or {}
    strategy_mode = str(dh.get("strategy_mode") or "").strip()
    if strategy_mode == "Survival":
        facts_only = report.get("facts_only") if isinstance(report.get("facts_only"), dict) else {}
        explicit_facts = " ".join(str(f) for f in (facts_only.get("explicit_facts") or [])).lower()
        survival_signals = [
            "без дохода", "no income", "долг", "debt",
            "срочно нужна", "urgent", "потери жилья", "housing risk",
            "финансовый дедлайн", "financial deadline",
            "быстрый доход как приоритет", "income urgency",
        ]
        if any(signal in explicit_facts for signal in survival_signals):
            return "Survival"
        readiness = dh.get("career_readiness") or {}
        urgency = str((readiness or {}).get("urgency", "")).lower()
        constraints = " ".join(str(c) for c in ((report.get("decision_layers") or {}).get("constraints") or [])).lower()
        if ("высок" in urgency or "high" in urgency) and any(s in constraints for s in survival_signals):
            return "Survival"
        return "Transition"
    if strategy_mode in ("Transition", "Growth"):
        return strategy_mode
    readiness = dh.get("career_readiness") or {}
    urgency = str((readiness or {}).get("urgency", "")).lower()
    if "сред" in urgency or "moder" in urgency or "medium" in urgency:
        return "Transition"
    return "Growth"


def build_meta(report: dict, user_name: str = "", profile_version: str = "") -> ReportMeta:
    from openai_client import REPORT_PIPELINE_VERSION  # noqa: PLC0415
    name = _safe_text(user_name, "Пользователь")
    country = _detect_country(report)
    mode = _detect_mode(report)
    created_at = datetime.now().strftime("%Y-%m-%d")
    pipeline_v = report.get("_pipeline_version") or REPORT_PIPELINE_VERSION
    return ReportMeta(
        user_name=name,
        country=country,
        mode=mode,
        created_at=created_at,
        profile_version=_safe_text(profile_version, "-") if str(profile_version or "").strip() else f"pipeline:{pipeline_v}",
    )


def _clean_fact_line(text: object) -> str:
    value = str(text or "").strip()
    if value.lower().startswith("ответ пользователя:"):
        value = value.split(":", 1)[1].strip()
    return value


def _build_story_echo(report: dict) -> str:
    facts_only = report.get("facts_only") if isinstance(report.get("facts_only"), dict) else {}
    explicit_facts = [_clean_fact_line(item) for item in (facts_only.get("explicit_facts") or [])]
    explicit_facts = [item for item in explicit_facts if item]
    inferences = [str(item).strip() for item in (facts_only.get("inferences") or []) if str(item).strip()]
    unknowns = [str(item).strip() for item in (facts_only.get("unknowns") or []) if str(item).strip()]

    facts: list[str] = []
    for item in explicit_facts:
        if item not in facts:
            facts.append(item)
        if len(facts) >= 4:
            break
    for item in inferences:
        if len(facts) >= 2:
            break
        if item not in facts:
            facts.append(item)
    if not facts:
        facts = ["Пока не хватает подтвержденных фактов из вашей истории. Это можно уточнить позже."]

    digital_human = report.get("digital_human") if isinstance(report.get("digital_human"), dict) else {}
    barriers = digital_human.get("barriers") if isinstance(digital_human.get("barriers"), dict) else {}

    problem = _safe_text(digital_human.get("main_barrier"), "данных недостаточно")
    if problem == "-":
        problem = "данных недостаточно"

    resource = _safe_text(digital_human.get("main_asset"), "данных недостаточно")
    if resource == "-":
        resource = "данных недостаточно"

    constraint_candidates = _list_items(barriers.get("external"))
    constraint = next((item for item in constraint_candidates if item != "-"), "")
    if not constraint and unknowns:
        constraint = unknowns[0]
    if not constraint:
        constraint = "данных недостаточно"

    facts_block = "\n".join(f"- {item}" for item in facts[:4])
    return (
        "Что я услышал в вашей истории\n\n"
        f"{facts_block}\n\n"
        f"Главная проблема:\n{problem}\n\n"
        f"Ресурс:\n{resource}\n\n"
        f"Ограничение:\n{constraint}"
    )


def _heard_facts(report: dict, limit: int = 4) -> list[str]:
    facts_only = report.get("facts_only") if isinstance(report.get("facts_only"), dict) else {}
    explicit_facts = [_clean_fact_line(item) for item in (facts_only.get("explicit_facts") or [])]
    explicit_facts = [item for item in explicit_facts if item]
    inferences = [str(item).strip() for item in (facts_only.get("inferences") or []) if str(item).strip()]
    combined: list[str] = []
    for item in explicit_facts + inferences:
        if item not in combined:
            combined.append(item)
        if len(combined) >= max(2, limit):
            break
    if not combined:
        return ["Подтвержденных фактов пока мало: уточните 2-4 ключевые детали опыта и текущей ситуации."]
    return combined[:limit]


def _main_request(report: dict) -> str:
    priorities = report.get("career_priorities") if isinstance(report.get("career_priorities"), list) else []
    cleaned_priorities = [str(item).strip() for item in priorities if str(item).strip()]
    if cleaned_priorities:
        return cleaned_priorities[0]
    decision = report.get("career_decision") if isinstance(report.get("career_decision"), dict) else {}
    route = _safe_text(decision.get("recommended_main_path"), "")
    if route:
        return f"Найти рабочий вход через маршрут: {route}"
    return "Запрос требует уточнения: какой результат важнее в ближайшие 30 дней"


def _strategy_bundle(report: dict) -> tuple[str, list[str]]:
    code = str(report.get("career_strategy") or "").strip()
    bundle = report.get(code) if code and isinstance(report.get(code), dict) else {}
    if not code or not isinstance(bundle, dict):
        return "", []

    if code == "fast_income":
        return (
            "Нужен доход в ближайшие 1–2 месяца",
            [
                _safe_text(bundle.get("goal_30_days")),
                f"Быстрые роли: {', '.join(_list_items(bundle.get('realistic_entry_roles'))[:4])}",
                f"Минимум для входа: {', '.join(_list_items(bundle.get('minimum_requirements'))[:5])}",
                f"План 7 дней: {', '.join(_list_items(bundle.get('application_plan_7_days'))[:3])}",
                f"Следующий шаг: {_safe_text(bundle.get('today_action', {}).get('action') if isinstance(bundle.get('today_action'), dict) else '')}",
            ],
        )
    if code == "upskill_for_profile":
        return (
            "Готов(а) добрать навыки 3–6 месяцев",
            [
                _safe_text(bundle.get("language_target")),
                f"Целевые роли: {', '.join(_list_items(bundle.get('target_roles_6_months'))[:4])}",
                f"План 12 недель: {', '.join(_list_items(bundle.get('checkpoints'))[:4])}",
                f"Портфолио / кейсы: {', '.join(_list_items(bundle.get('portfolio_or_case_plan'))[:3])}",
            ],
        )
    if code == "long_transition":
        return (
            "Готов(а) вложиться в переобучение и смену траектории",
            [
                _safe_text(bundle.get("goal_30_days")),
                f"Долгий горизонт: {', '.join(_list_items(bundle.get('target_roles_6_months'))[:4])}",
                f"Сегодня: {_safe_text(bundle.get('today_action', {}).get('action') if isinstance(bundle.get('today_action'), dict) else '')}",
            ],
        )
    if code == "need_decision":
        return (
            "Не уверен(а), помоги выбрать",
            [
                _safe_text(bundle.get("message")),
                f"Предварительный маршрут: {_safe_text(bundle.get('preliminary_route'))}",
                f"Что уточнить: {', '.join(_list_items(bundle.get('missing_fields'))[:6]) or 'данных недостаточно'}",
            ],
        )
    return "", []


def _unknowns_list(report: dict) -> list[str]:
    facts_only = report.get("facts_only") if isinstance(report.get("facts_only"), dict) else {}
    unknowns = [str(item).strip() for item in (facts_only.get("unknowns") or []) if str(item).strip()]
    return unknowns[:6]


def _decatastrophize_language(text: object) -> str:
    value = str(text or "").strip()
    if not value:
        return "данных недостаточно"
    replacements = {
        "Ваша идеальная профессия": "Возможный маршрут",
        "Вы точно должны стать": "Предварительная гипотеза",
        "Ваше настоящее предназначение": "Потребуется проверить",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value


def _route_evidence_blocks(report: dict) -> list[dict[str, object]]:
    raw = report.get("route_evidence_blocks") if isinstance(report.get("route_evidence_blocks"), list) else []
    blocks: list[dict[str, object]] = []
    allowed_roles = {"primary", "transition", "quick", "emergency"}

    def _normalize_row(row: dict, fallback_income_role: str) -> dict[str, object]:
        income_role = str(row.get("income_role") or fallback_income_role).strip().lower()
        if income_role not in allowed_roles:
            income_role = fallback_income_role
        return {
            "route": _decatastrophize_language(row.get("route") or "Возможный маршрут"),
            "why_it_fits": _list_items(row.get("why_it_fits")),
            "evidence_from_user": _list_items(row.get("evidence_from_user")),
            "missing_competencies": _list_items(row.get("missing_competencies")),
            "entry_level": _safe_text(row.get("entry_level"), "Потребуется проверить"),
            "income_role": income_role,
            "risks": _list_items(row.get("risks")),
            "what_may_disprove_this_route": _list_items(row.get("what_may_disprove_this_route")),
        }

    for row in raw:
        if not isinstance(row, dict):
            continue
        fallback = "primary" if len(blocks) == 0 else "transition" if len(blocks) == 1 else "quick" if len(blocks) == 2 else "emergency"
        blocks.append(_normalize_row(row, fallback))

    if blocks:
        return blocks

    market = report.get("market_analysis") if isinstance(report.get("market_analysis"), list) else []
    facts = _heard_facts(report, limit=4)
    fallback_roles = ["primary", "transition", "quick", "emergency"]
    for idx, role in enumerate(fallback_roles):
        row = market[idx] if idx < len(market) and isinstance(market[idx], dict) else {}
        route_name = _safe_text(row.get("profession"), "Возможный маршрут")
        risks = [
            _safe_text(row.get("competition"), "Потребуется проверить уровень конкуренции"),
            "Потребуется проверить стабильность входа на локальном рынке",
        ]
        blocks.append(
            {
                "route": _decatastrophize_language(route_name),
                "why_it_fits": [
                    _safe_text(row.get("profile_match_reason"), "Предварительная гипотеза по совпадению опыта."),
                ],
                "evidence_from_user": facts[:3],
                "missing_competencies": _list_items(row.get("requirements"))[:4],
                "entry_level": _safe_text(row.get("entry_speed"), "Потребуется проверить"),
                "income_role": role,
                "risks": risks[:3],
                "what_may_disprove_this_route": [
                    "Новые факты о legal access или языковом уровне",
                    "Прямой отказ пользователя от функции маршрута",
                ],
            }
        )
    return blocks


def _psych_social_recommendation(report: dict) -> tuple[list[str], list[str]]:
    digital_human = report.get("digital_human") if isinstance(report.get("digital_human"), dict) else {}
    barriers = digital_human.get("barriers") if isinstance(digital_human.get("barriers"), dict) else {}
    social_integration = report.get("social_integration") if isinstance(report.get("social_integration"), dict) else {}

    internal = _list_items(barriers.get("internal"))[:3]
    external = _list_items(barriers.get("external"))[:3]
    fears = _list_items((digital_human.get("psychological_profile") or {}).get("dominant_fears"))[:3]
    communities = _list_items(social_integration.get("communities"))[:3]

    self_help: list[str] = []
    if any(item != "-" for item in internal):
        self_help.append("Действовать через микрошаг 5-15 минут в день, без попытки решить всё сразу.")
    if any(item != "-" for item in external):
        self_help.append("Разделить барьеры на управляемые (сегодня) и внешние (план на неделю).")
    if any(item != "-" for item in communities):
        self_help.append("Добавить один социальный контакт в неделю: сообщество, чат или профильный канал.")
    if not self_help:
        self_help.append("Пока данных мало: начните с одного короткого действия и зафиксируйте результат.")

    specialist: list[str] = []
    signal_blob = " ".join([" ".join(fears), " ".join(internal)]).lower().replace("ё", "е")
    if any(token in signal_blob for token in ["страх", "пан", "стыд", "не могу", "нет сил", "хаос", "трев"]):
        specialist.append("Если перегруз и тревога держатся неделями, полезно разобрать это с психологом/консультантом по адаптации.")
    specialist.append("Если маршрут понятен, но нет движения, полезен разбор со специалистом: барьер, темп, корректировка шага.")
    specialist.append("Это не медицинское заключение: рекомендация нужна только для опоры и ускорения адаптации.")

    return self_help[:3], specialist[:3]


def _first_step_buttons(report: dict) -> list[str]:
    buttons = report.get("first_step_buttons") if isinstance(report.get("first_step_buttons"), list) else []
    cleaned = [str(item).strip() for item in buttons if str(item).strip()]
    if cleaned:
        return cleaned[:5]

    if str(report.get("profile_domain") or "").strip() == "construction_engineering_cost_estimation":
        return [
            "Сделал",
            "Слишком сложно",
            "Сделать проще",
            "Помоги составить таблицу",
            "Хочу примеры запросов",
        ]

    return ["Сделал", "Слишком сложно", "Сделать проще", "Другой шаг"]


def _guidance_items(value: object, text_key: str, detail_key: str = "") -> list[str]:
    """Render structured continuation items without leaking internal enum values."""
    if not isinstance(value, list):
        return []
    result: list[str] = []
    for item in value:
        if isinstance(item, dict):
            text = str(item.get(text_key) or "").strip()
            detail = str(item.get(detail_key) or "").strip() if detail_key else ""
            if text:
                result.append(f"{text} — {detail}" if detail else text)
        elif str(item).strip():
            result.append(str(item).strip())
    return result


def _next_step_guidance(report: dict) -> dict[str, object]:
    """Return model guidance with safe, useful fallbacks for legacy reports."""
    raw = report.get("next_step_guidance")
    guidance = dict(raw) if isinstance(raw, dict) else {}
    decision = report.get("career_decision") if isinstance(report.get("career_decision"), dict) else {}
    facts = report.get("facts_only") if isinstance(report.get("facts_only"), dict) else {}
    action = report.get("action_plan") if isinstance(report.get("action_plan"), dict) else {}
    today = action.get("today") if isinstance(action.get("today"), dict) else {}
    route = _decision_route_title(decision)
    unknowns = [str(x).strip() for x in facts.get("unknowns", []) if str(x).strip()]

    if not guidance.get("main_risks"):
        guidance["main_risks"] = [{
        "risk": _safe_text((report.get("digital_human") or {}).get("main_risk"), f"Проверять маршрут «{route}» только по названию роли."),
        "evidence": _safe_text(decision.get("why_this_path"), "Маршрут пока является рабочей гипотезой."),
        "consequence": "До вложений в обучение нужно сверить гипотезу с реальными требованиями вакансий.",
        }]
    if not guidance.get("checks_before_decision"):
        guidance["checks_before_decision"] = [
        {"check": item, "why_it_matters": "Ответ может изменить маршрут, срок входа или допустимый уровень дохода."}
        for item in unknowns[:4]
        ] or [{"check": f"Требования 10 вакансий по маршруту «{route}»", "why_it_matters": "Это покажет доступность входа до затрат на обучение."}]
    if not guidance.get("self_service_actions"):
        guidance["self_service_actions"] = [
            {
                "action": _safe_text(today.get("action"), f"Собрать 10 вакансий по маршруту «{route}»."),
                "result": _safe_text(today.get("result"), "Список повторяющихся требований."),
            },
            {"action": "Выписать требования, которые повторяются минимум в трёх вакансиях", "result": "Приоритетный список пробелов без покупки обучения вслепую"},
        ]
    elif isinstance(guidance.get("self_service_actions"), list) and len(guidance["self_service_actions"]) < 2:
        guidance["self_service_actions"] = list(guidance["self_service_actions"]) + [{
            "action": "Сверить выбранный маршрут минимум с пятью реальными вакансиями",
            "result": "Список подтверждённых требований и вопросов, которые ещё нужно проверить",
        }]
    if not guidance.get("support_accelerators"):
        guidance["support_accelerators"] = [
        {"task": "Сопоставить реальные вакансии на одной системе критериев", "result": "Подтверждённый основной и запасной маршрут", "format": "both"},
        {"task": "Адаптировать CV и отклики под повторяющиеся требования", "result": "Готовый комплект для рыночного теста", "format": "ai"},
        ]
    if not guidance.get("decision_level"):
        guidance["decision_level"] = {
        "known": f"Рабочий основной маршрут — {route}.",
        "next_confirmation": "рынок → доступность входа → деньги",
        "decision_after": "После проверки можно решать вопрос об обучении и переходе.",
        }
    route_count = len(report.get("selected_routes", [])) if isinstance(report.get("selected_routes"), list) else 0
    route_count = route_count or len(report.get("route_evidence_blocks", [])) if isinstance(report.get("route_evidence_blocks"), list) else route_count
    high_stakes = any(token in " ".join(_guidance_items(guidance.get("main_risks"), "risk")).lower() for token in ("доход", "деньг", "обуч", "уволь", "финанс"))
    if unknowns:
        cta_type, cta_title, why_now = "career_chat", "Продолжить разбор в карьерном чате", "Сначала нужно собрать недостающие факты."
    elif 2 <= route_count <= 4 and high_stakes:
        cta_type, cta_title, why_now = "career_consultant", "Разобрать решение с карьерным консультантом", "Есть несколько реалистичных маршрутов с существенными компромиссами."
    else:
        cta_type, cta_title, why_now = "job_search_support", "Перейти к сопровождению поиска работы", "Маршрут уже можно проверять откликами."
    if not guidance.get("primary_cta"):
        guidance["primary_cta"] = {
        "type": cta_type,
        "title": cta_title,
        "why_now": why_now,
        "outcomes": ["Разобрать вакансии", "Уточнить обязательные требования", "Обновить стратегию по результатам"],
        }
    if not guidance.get("first_chat_task"):
        guidance["first_chat_task"] = {
        "action": _safe_text(today.get("action"), f"Найдите 5 вакансий по маршруту «{route}»."),
        "volume": "5 вакансий",
        "result_to_send": "Ссылки или тексты вакансий",
        "assistant_response": "Я сравню требования, отмечу повторяющиеся навыки и предложу следующий шаг.",
        }
    if not guidance.get("human_escalation_triggers"):
        guidance["human_escalation_triggers"] = [
            "два близких маршрута дают разные финансовые последствия",
            "рыночные данные противоречат друг другу или тестовые отклики долго не дают результата",
            "нужна персональная обратная связь по интервью, портфолио или переговорам о зарплате",
        ]
    report["next_step_guidance"] = guidance
    return guidance


def ensure_next_step_guidance(report: dict) -> dict[str, object]:
    """Persist a complete continuation contract before report delivery/state storage."""
    return _next_step_guidance(report)


def _guidance_text(report: dict) -> str:
    guidance = _next_step_guidance(report)
    risks = _guidance_items(guidance.get("main_risks"), "risk", "consequence")[:3]
    checks = _guidance_items(guidance.get("checks_before_decision"), "check", "why_it_matters")[:5]
    own = _guidance_items(guidance.get("self_service_actions"), "action", "result")[:5]
    support = _guidance_items(guidance.get("support_accelerators"), "task", "result")[:6]
    level = guidance.get("decision_level") if isinstance(guidance.get("decision_level"), dict) else {}
    cta = guidance.get("primary_cta") if isinstance(guidance.get("primary_cta"), dict) else {}
    task = guidance.get("first_chat_task") if isinstance(guidance.get("first_chat_task"), dict) else {}
    alternative = guidance.get("alternative_cta") if isinstance(guidance.get("alternative_cta"), dict) else {}
    outcomes = [str(x).strip() for x in cta.get("outcomes", []) if str(x).strip()]
    block = lambda items: "\n".join(f"- {x}" for x in items) or "- Критичных пунктов не выявлено."
    sections = [
        f"Где сейчас главный риск ошибиться\n{block(risks)}",
        f"Что ещё нужно проверить\n{block(checks)}",
        f"Что вы можете сделать самостоятельно\n{block(own)}",
        f"Где сопровождение может ускорить переход\n{block(support)}",
        "Ваш следующий уровень решения\n" + " ".join(str(level.get(k) or "").strip() for k in ("known", "next_confirmation", "decision_after")).strip(),
        f"Следующий шаг — {_safe_text(cta.get('title'))}\n{_safe_text(cta.get('why_now'))}\n{block(outcomes)}",
        f"Первое действие в чате\n{_safe_text(task.get('action'))} ({_safe_text(task.get('volume'))}). "
        f"Пришлите: {_safe_text(task.get('result_to_send'))}. {_safe_text(task.get('assistant_response'))}",
        "Когда особенно полезен живой консультант\n" + block([str(x).strip() for x in guidance.get("human_escalation_triggers", []) if str(x).strip()]),
        (f"Альтернативный формат — {_safe_text(alternative.get('title'))}\n{_safe_text(alternative.get('why_now'))}" if alternative else ""),
    ]
    return "\n\n".join(section for section in sections if section)


def build_telegram_summary(report: dict) -> str:
    from services.market_strategy import humanize_internal_values
    report = humanize_internal_values(report)
    digital_human = report.get("digital_human", {}) if isinstance(report.get("digital_human"), dict) else {}
    decision = report.get("career_decision", {}) if isinstance(report.get("career_decision"), dict) else {}
    action_plan = report.get("action_plan", {}) if isinstance(report.get("action_plan"), dict) else {}
    today = action_plan.get("today", {}) if isinstance(action_plan.get("today"), dict) else {}
    weekly = report.get("weekly_plan", []) if isinstance(report.get("weekly_plan"), list) else []
    development = report.get("development_map", {}) if isinstance(report.get("development_map"), dict) else {}
    not_reset = report.get("what_not_reset", []) if isinstance(report.get("what_not_reset"), list) else []
    not_reset_block = "\n".join(f"- {item}" for item in _list_items(not_reset)[:4])
    energy_sources = report.get("energy_sources", []) if isinstance(report.get("energy_sources"), list) else []
    energy_block = "\n".join(f"- {item}" for item in _list_items(energy_sources)[:4])
    career_priorities = report.get("career_priorities", []) if isinstance(report.get("career_priorities"), list) else []
    priorities_block = "\n".join(f"- {item}" for item in _list_items(career_priorities)[:4])
    competency_signals = report.get("competency_signals", []) if isinstance(report.get("competency_signals"), list) else []
    competency_block = "\n".join(f"- {item}" for item in _list_items(competency_signals)[:5])
    weaknesses = []
    barriers = (digital_human.get("barriers") or {}) if isinstance(digital_human.get("barriers"), dict) else {}
    weaknesses.extend(_list_items(barriers.get("internal"))[:2])
    weaknesses.extend(_list_items(barriers.get("external"))[:2])
    weaknesses_block = "\n".join(f"- {item}" for item in weaknesses[:4]) if weaknesses else "- данных недостаточно"
    facts_only = report.get("facts_only") if isinstance(report.get("facts_only"), dict) else {}
    unknowns = [str(item).strip() for item in (facts_only.get("unknowns") or []) if str(item).strip()]
    unknowns_block = "\n".join(f"- {item}" for item in unknowns[:3]) if unknowns else "- критичных неизвестных сейчас нет"
    resource_level = _resource_human_message(report.get("resource_level"))
    integration_level = _integration_human_message(report.get("integration_level"))
    story_echo = _build_story_echo(report)
    professional_core = _professional_core_summary(report)
    market = report.get("market_analysis", []) if isinstance(report.get("market_analysis"), list) else []
    resume_analysis = report.get("resume_analysis", {}) if isinstance(report.get("resume_analysis"), dict) else {}
    resume_status = "есть" if resume_analysis else "нет"
    route_lines: list[str] = []
    for idx, item in enumerate(market[:3], start=1):
        if not isinstance(item, dict):
            continue
        route_lines.append(
            f"{idx}. {_safe_text(item.get('profession'))}: {_safe_text(item.get('salary_range'))}; {_safe_text(item.get('fit_percent'))}%"
        )
    route_block = "\n".join(f"- {item}" for item in route_lines) if route_lines else "- данных недостаточно"
    weekly_block = "\n".join(f"- День {item.get('day', '-')}: {_safe_text(item.get('task'))}" for item in weekly[:4] if isinstance(item, dict)) or "- данных недостаточно"
    first_month = development.get("first_month", []) if isinstance(development.get("first_month"), list) else []
    month_block = "\n".join(
        f"- Неделя {row.get('week', '-')}: {_safe_text(row.get('focus'))} — {_safe_text(row.get('output'))}"
        for row in first_month[:4]
        if isinstance(row, dict)
    ) or "- данных недостаточно"
    strategy_title, strategy_lines = _strategy_bundle(report)
    strategy_block = "\n".join(f"- {line}" for line in strategy_lines) if strategy_lines else "- данных недостаточно"
    route_evidence = _route_evidence_blocks(report)

    income_label = {
        "primary": "Основной маршрут",
        "transition": "Переходный маршрут",
        "quick": "Быстрый доход",
        "emergency": "Аварийный вариант",
    }
    evidence_lines: list[str] = []
    for item in route_evidence[:4]:
        if not isinstance(item, dict):
            continue
        label = income_label.get(str(item.get("income_role") or ""), "Возможный маршрут")
        evidence_lines.append(
            f"- {label}: {_safe_text(item.get('route'))}; доказательства: {', '.join(_list_items(item.get('evidence_from_user'))[:2])}; "
            f"проверка: {', '.join(_list_items(item.get('what_may_disprove_this_route'))[:2])}"
        )
    evidence_block = "\n".join(evidence_lines) if evidence_lines else "- данных недостаточно"

    function_lines = [f"- {item}" for item in _list_items(competency_signals)[:5] if str(item).strip() and str(item).strip() != "-"]
    function_block = "\n".join(function_lines) if function_lines else "- данных недостаточно"

    level_lines = []
    for row in report.get("function_levels", []) if isinstance(report.get("function_levels"), list) else []:
        if not isinstance(row, dict):
            continue
        level_lines.append(f"- {_safe_text(row.get('function'))}: {_safe_text(row.get('level'))}")
    if not level_lines:
        level_lines.append("- Предварительная гипотеза: уровень по функциям нужно уточнить по доказательствам.")
    levels_block = "\n".join(level_lines)

    explicit_refusals = report.get("explicit_refusals") if isinstance(report.get("explicit_refusals"), list) else []
    refusals_block = "\n".join(f"- {str(item).strip()}" for item in explicit_refusals if str(item).strip())
    if not refusals_block:
        refusals_block = "- Прямые отказы не зафиксированы: потребуется проверить."

    cannot_claim_block = "\n".join(f"- {item}" for item in unknowns[:4]) if unknowns else "- Категоричные выводы делать нельзя: данных недостаточно."
    constraints_block = "\n".join(f"- {item}" for item in weaknesses[:5] if str(item).strip()) or "- данных недостаточно"
    main_route = _decision_route_title(decision)
    backup_route = _safe_text(decision.get("backup_path"), "Предварительная гипотеза")

    summary = [
        story_echo,
        "",
        "Ваш NextYou отчёт",
        "",
        f"1. Как мы поняли вашу ситуацию:\n{story_echo}",
        f"2. Ваше профессиональное ядро:\n{professional_core}",
        f"3. Подтверждённые функции:\n{function_block}",
        f"4. Уровень по каждой функции:\n{levels_block}",
        f"5. Что нельзя утверждать на основании имеющихся данных:\n{cannot_claim_block}",
        f"6. Что вы хотите сохранить:\n{not_reset_block}",
        f"7. От чего вы хотите уйти:\n{refusals_block}",
        f"8. Основной маршрут:\nВозможный маршрут: {main_route}",
        f"9. Переходный маршрут:\nПредварительная гипотеза: {backup_route}",
        f"10. Быстрый доход:\n{_safe_text(digital_human.get('fastest_path_to_income'), 'Потребуется проверить конкретную роль быстрого дохода.')}",
        f"11. Аварийный вариант:\n{_safe_text(decision.get('avoid_for_now'), 'Потребуется проверить аварийный вариант в зависимости от ограничений.')}",
        f"12. Требования и ограничения:\n{constraints_block}",
        f"13. Неопределённости:\n{unknowns_block}",
        f"Маршрутные доказательства:\n{evidence_block}",
        f"14. Первый шаг до 15 минут:\n{_safe_text(today.get('action'))}",
        f"15. План на неделю и месяц:\n{weekly_block}\n{month_block}",
        _guidance_text(report),
        f"Стратегия выбора:\n{strategy_title or 'данных недостаточно'}\n{strategy_block}",
        f"Кто вы сейчас:\n{_safe_text(digital_human.get('current_state'))}",
        f"Что не обнулилось:\n{not_reset_block}",
        f"Ваше профессиональное ядро:\n{professional_core}",
        f"Сильные стороны и опоры:\n{energy_block}",
        f"Ограничения и неизвестные:\n{weaknesses_block}\n\nЧто уточнить:\n{unknowns_block}",
        f"Устойчивость в период изменений:\n{resource_level}",
        f"Ресурс и рабочий темп:\n{resource_level}",
        f"Интеграция в новой стране:\n{integration_level}",
        f"Состояние интеграции:\n{integration_level}",
        f"Сравнение маршрутов:\n{route_block}",
        f"Выбранный маршрут и первый шаг:\n{_safe_text(decision.get('recommended_main_path'))}\n{_safe_text(today.get('action'))}",
        f"План на 30 дней:\n{weekly_block}\n{month_block}",
        f"10. Анализ резюме:\nАнализ CV: {resume_status}",
        f"11. Что может быть не так:\n{weaknesses_block}",
        f"Источники энергии:\n{energy_block}",
        f"Карьерные приоритеты:\n{priorities_block}",
        f"STAR-компетенции:\n{competency_block}",
        f"Что делать самому:\n{_safe_text(today.get('action'))}",
        f"Что подумать со специалистом:\n{_safe_text(decision.get('why_this_path'))}",
        f"Главный риск:\n{_safe_text(digital_human.get('main_risk'))}",
        f"Рекомендуемый маршрут:\n{main_route}",
        f"Почему он предложен:\n{_safe_text(decision.get('why_this_path'))}",
        f"Запасной маршрут:\n{backup_route}",
        f"Маленький первый шаг:\n{_safe_text(today.get('action'))}",
        "Что хотите сделать дальше?",
    ]
    return "\n\n".join(summary)


def build_offer_text() -> str:
    return (
        "Мы можем не просто дать карту, а помочь пройти первые шаги до реальных вакансий.\n\n"
        "Что входит:\n"
        "- уточнение маршрута;\n"
        "- анализ CV;\n"
        "- подбор первых вакансий;\n"
        "- подготовка откликов;\n"
        "- разбор страхов и барьеров;\n"
        "- план действий на неделю;\n"
        "- корректировка действий.\n\n"
        "Мы уже знаем вашу ситуацию. Теперь можем помочь дойти до первых вакансий."
    )


def _decision_route_title(decision: dict) -> str:
    main_route = decision.get("main_route") if isinstance(decision.get("main_route"), dict) else None
    if main_route:
        title = str(main_route.get("title") or main_route.get("name") or "").strip()
        if title:
            return title
    return str(decision.get("recommended_main_path") or decision.get("main_route") or decision.get("route") or "Возможный маршрут").strip() or "Возможный маршрут"


def _decision_country_label(report: dict, decision: dict) -> tuple[str, str]:
    snapshot = report.get("profile_snapshot") if isinstance(report.get("profile_snapshot"), dict) else {}
    country_name = str(decision.get("country_name") or snapshot.get("country_name") or report.get("country_name") or "").strip()
    city = str(decision.get("city") or snapshot.get("city") or report.get("city") or "").strip()
    if not country_name:
        country_name = str((snapshot.get("route_context") or {}).get("country") or "").strip() or "Не уточнено"
    if not city:
        city = str((snapshot.get("route_context") or {}).get("city") or "").strip()
    return country_name or "Не уточнено", city or "Не уточнено"


def render_report_html(report: dict, meta: ReportMeta) -> str:
    from services.market_strategy import humanize_internal_values
    report = humanize_internal_values(report)
    digital_human = report.get("digital_human", {}) if isinstance(report.get("digital_human"), dict) else {}
    decision = report.get("career_decision", {}) if isinstance(report.get("career_decision"), dict) else {}
    country_name, city = _decision_country_label(report, decision)
    market = report.get("market_analysis", []) if isinstance(report.get("market_analysis"), list) else []
    recommendations = report.get("career_recommendations", []) if isinstance(report.get("career_recommendations"), list) else []
    real_solutions = report.get("real_solutions", []) if isinstance(report.get("real_solutions"), list) else []
    translation = report.get("career_translation", []) if isinstance(report.get("career_translation"), list) else []
    bridges = report.get("career_bridges", []) if isinstance(report.get("career_bridges"), list) else []
    not_reset = report.get("what_not_reset", []) if isinstance(report.get("what_not_reset"), list) else []
    experience_layers = report.get("experience_layers", []) if isinstance(report.get("experience_layers"), list) else []
    career_barriers = report.get("career_barriers", []) if isinstance(report.get("career_barriers"), list) else []
    barrier_landscape = report.get("barrier_landscape", {}) if isinstance(report.get("barrier_landscape"), dict) else {}
    action_plan = report.get("action_plan", {}) if isinstance(report.get("action_plan"), dict) else {}
    weekly = report.get("weekly_plan", []) if isinstance(report.get("weekly_plan"), list) else []
    development = report.get("development_map", {}) if isinstance(report.get("development_map"), dict) else {}
    social_integration = report.get("social_integration", {}) if isinstance(report.get("social_integration"), dict) else {}
    energy_sources = report.get("energy_sources", []) if isinstance(report.get("energy_sources"), list) else []
    career_priorities = report.get("career_priorities", []) if isinstance(report.get("career_priorities"), list) else []
    competency_signals = report.get("competency_signals", []) if isinstance(report.get("competency_signals"), list) else []
    resume_analysis = report.get("resume_analysis", {}) if isinstance(report.get("resume_analysis"), dict) else {}
    resource_level = _resource_human_message(report.get("resource_level"))
    integration_level = _integration_human_message(report.get("integration_level"))
    closing_message = _safe_text(
        report.get("closing_message"),
        "Это работа, а не испытание. Один проверяемый шаг — и карта начнёт двигаться.",
    )
    strengths_for_closing = list(dict.fromkeys(
        [item for item in _list_items(not_reset)[:3] if item != "-"] +
        [item for item in _list_items(digital_human.get("hidden_strengths"))[:3] if item != "-"]
    ))[:4]
    first_step_today = _safe_text(
        (report.get("action_plan") or {}).get("today", {}).get("action") if isinstance((report.get("action_plan") or {}).get("today"), dict) else None,
        "Собрать 10 вакансий по выбранному маршруту и выписать повторяющиеся требования.",
    )
    next_step_week = _safe_text(
        ((report.get("action_plan") or {}).get("this_week") or [""])[0] if isinstance((report.get("action_plan") or {}).get("this_week"), list) else "",
        "Подготовить CV под выбранный маршрут и отправить 3-5 тестовых отклика.",
    )
    specialist_hints, _ = _psych_social_recommendation(report)
    weekly_signals = [
        _safe_text(item.get("task"))
        for item in weekly[:3]
        if isinstance(item, dict)
    ]
    strengths_items = list(dict.fromkeys(_list_items(not_reset)[:4] + _list_items(digital_human.get("hidden_strengths"))[:4]))[:6]
    barriers_obj = digital_human.get("barriers") if isinstance(digital_human.get("barriers"), dict) else {}
    weaknesses_items = list(dict.fromkeys(_list_items(barriers_obj.get("internal"))[:3] + _list_items(barriers_obj.get("external"))[:3]))[:6]
    blind_spots = [
        _safe_text(digital_human.get("main_risk"), ""),
        _safe_text(digital_human.get("main_barrier"), ""),
        _safe_text((decision or {}).get("avoid_for_now"), ""),
        _safe_text((barrier_landscape or {}).get("behavioral_risk"), ""),
    ]
    blind_spots = [item for item in blind_spots if item and item != "-"][:4]
    important_consider = [
        f"Ресурс: {resource_level.splitlines()[0]}",
        f"Интеграция: {integration_level.splitlines()[0]}",
        f"Приоритеты: {', '.join(_list_items(career_priorities)[:3])}",
    ]
    strategy_title, strategy_lines = _strategy_bundle(report)
    strategy_html = (
        f"<div class='card'><h3>Стратегия выбора</h3><p><b>{escape(strategy_title)}</b></p><ul>"
        + ''.join(f"<li>{escape(line)}</li>" for line in strategy_lines)
        + "</ul></div>"
    ) if strategy_lines else "<div class='card'><h3>Стратегия выбора</h3><p>Данных недостаточно.</p></div>"
    market_questions = []
    for item in market[:6]:
        if not isinstance(item, dict):
            continue
        for req in _list_items(item.get("requirements"))[:4]:
            if req not in market_questions and req != "-":
                market_questions.append(req)
    unknowns = _unknowns_list(report)
    heard_facts = _heard_facts(report, limit=4)
    main_request = _main_request(report)
    primary_constraint = next((item for item in weaknesses_items if item != "-"), "данных недостаточно")
    self_help_points, specialist_points = _psych_social_recommendation(report)
    scenario_labels = ["Безопасный", "Основной", "Амбициозный"]
    route_evidence = _route_evidence_blocks(report)

    function_lines = "".join(f"<li>{escape(item)}</li>" for item in _list_items(competency_signals)[:6]) or "<li>Данных недостаточно.</li>"
    function_levels = report.get("function_levels") if isinstance(report.get("function_levels"), list) else []
    level_lines = "".join(
        f"<li>{escape(_safe_text(row.get('function')))}: {escape(_safe_text(row.get('level')))}</li>"
        for row in function_levels[:8]
        if isinstance(row, dict)
    ) or "<li>Предварительная гипотеза: уровень по функциям потребует проверки.</li>"
    explicit_refusals = report.get("explicit_refusals") if isinstance(report.get("explicit_refusals"), list) else []
    refusals_html = "".join(f"<li>{escape(str(item))}</li>" for item in explicit_refusals[:6] if str(item).strip()) or "<li>Прямые отказы не зафиксированы: потребуется проверить.</li>"
    unknowns_html = "".join(f"<li>{escape(item)}</li>" for item in unknowns[:6]) or "<li>Категоричные выводы делать нельзя: данных недостаточно.</li>"

    route_label = {
        "primary": "Основной маршрут",
        "transition": "Переходный маршрут",
        "quick": "Быстрый доход",
        "emergency": "Аварийный вариант",
    }
    route_evidence_html = "".join(
        (
            "<div class='card'>"
            f"<h3>{escape(route_label.get(str(item.get('income_role') or ''), 'Возможный маршрут'))}: {escape(_safe_text(item.get('route')))}</h3>"
            f"<p><b>Почему подходит:</b> {escape(', '.join(_list_items(item.get('why_it_fits'))[:4]))}</p>"
            f"<p><b>Доказательства пользователя:</b> {escape(', '.join(_list_items(item.get('evidence_from_user'))[:4]))}</p>"
            f"<p><b>Чего не хватает:</b> {escape(', '.join(_list_items(item.get('missing_competencies'))[:4]))}</p>"
            f"<p><b>Уровень входа:</b> {escape(_safe_text(item.get('entry_level')))}</p>"
            f"<p><b>Риски:</b> {escape(', '.join(_list_items(item.get('risks'))[:4]))}</p>"
            f"<p><b>Что может опровергнуть маршрут:</b> {escape(', '.join(_list_items(item.get('what_may_disprove_this_route'))[:4]))}</p>"
            "</div>"
        )
        for item in route_evidence[:4]
        if isinstance(item, dict)
    ) or "<div class='card'><p>Данных недостаточно для доказательного блока маршрутов.</p></div>"

    possibilities = []
    labels = ["Быстрый доход", "Основной маршрут", "Долгосрочное развитие"]
    for idx, item in enumerate(market[:3]):
        if not isinstance(item, dict):
            continue
        possibilities.append(
            f"""
            <div class='card'>
              <h3>{labels[idx]}: {escape(_safe_text(item.get('profession')))}</h3>
              <ul>
                <li><b>Соответствие:</b> {escape(str(item.get('fit_percent', '-')))}%</li>
                <li><b>Скорость входа:</b> {escape(_safe_text(item.get('entry_speed')))}</li>
                <li><b>Риск:</b> {escape(_safe_text(item.get('competition')))}</li>
                <li><b>Доход:</b> {escape(_safe_text(item.get('salary_range')))}</li>
                <li><b>Что подтянуть:</b> {escape(', '.join(_list_items(item.get('requirements'))[:5]))}</li>
              </ul>
            </div>
            """
        )

    recommendations_html = []
    for item in recommendations[:4]:
        if not isinstance(item, dict):
            continue
        recommendations_html.append(
            f"""
            <div class='card'>
                <h3>{escape(_safe_text(item.get('title')))}</h3>
                <ul>
                    <li><b>Соответствие:</b> {escape(str(item.get('match_percent', '-')))}%</li>
                    <li><b>Почему подходит:</b> {escape(_safe_text(item.get('why_fit')))}</li>
                    <li><b>Риски:</b> {escape(', '.join(_list_items(item.get('risks'))[:4]))}</li>
                    <li><b>Срок входа:</b> {escape(_safe_text(item.get('entry_timeline')))}</li>
                    <li><b>Доход:</b> {escape(_safe_text(item.get('income_range')))}</li>
                </ul>
            </div>
            """
        )

    solutions_html = []
    for idx, item in enumerate(real_solutions[:3]):
        if not isinstance(item, dict):
            continue
        solutions_html.append(
            f"""
            <div class='card'>
                <h3>{escape(scenario_labels[idx])}: {escape(_safe_text(item.get('title')))}</h3>
                <ul>
                    <li><b>Уровень рекомендации:</b> {escape(_safe_text(item.get('recommendation_level')))}</li>
                    <li><b>Вероятность успеха:</b> {escape(_safe_text(item.get('success_probability')))}</li>
                    <li><b>Срок:</b> {escape(_safe_text(item.get('timeline')))}</li>
                    <li><b>Почему:</b> {escape(_safe_text(item.get('why')))}</li>
                    <li><b>Первый шаг:</b> {escape(_safe_text(item.get('first_step')))}</li>
                </ul>
            </div>
            """
        )

    translation_html = []
    for item in translation[:6]:
        if not isinstance(item, dict):
            continue
        translation_html.append(
            f"""
            <div class='card'>
                <h3>Перевод опыта</h3>
                <p><b>Было:</b> {escape(_safe_text(item.get('source_experience')))}</p>
                <p><b>На рынке называется:</b> {escape(_safe_text(item.get('market_term')))}</p>
                <p><b>Подходящие роли:</b> {escape(', '.join(_list_items(item.get('suitable_roles'))[:4]))}</p>
            </div>
            """
        )

    bridges_html = []
    for item in bridges[:5]:
        if not isinstance(item, dict):
            continue
        bridges_html.append(
            f"""
            <div class='card'>
                <h3>{escape(_safe_text(item.get('role')))}</h3>
                <p><b>Почему это мост:</b> {escape(_safe_text(item.get('why_bridge')))}</p>
                <p><b>Первый тест рынка:</b> {escape(_safe_text(item.get('first_market_test')))}</p>
            </div>
            """
        )

    barriers_html = []
    for item in career_barriers[:5]:
        if not isinstance(item, dict):
            continue
        barriers_html.append(
            f"""
            <div class='card'>
                <h3>{escape(_safe_text(item.get('barrier')))}</h3>
                <ul>
                    <li><b>Сила влияния:</b> {escape(str(item.get('severity', '-')))} / 100</li>
                    <li><b>Что это значит:</b> {escape(_safe_text(item.get('mechanism')))}</li>
                    <li><b>Навык для компенсации:</b> {escape(_safe_text(item.get('recommended_skill')))}</li>
                    <li><b>Первое упражнение:</b> {escape(_safe_text(item.get('first_exercise')))}</li>
                </ul>
            </div>
            """
        )

    today = action_plan.get("today", {}) if isinstance(action_plan.get("today"), dict) else {}
    this_week = "".join(f"<li>{escape(str(step))}</li>" for step in _list_items(action_plan.get("this_week")))
    this_month = "".join(f"<li>{escape(str(step))}</li>" for step in _list_items(action_plan.get("this_month")))
    first_month = development.get("first_month", []) if isinstance(development.get("first_month"), list) else []
    month_rows = "".join(
        f"<li><b>Неделя {escape(str(row.get('week', '-')))}:</b> {escape(_safe_text(row.get('focus')))} — {escape(_safe_text(row.get('output')))}</li>"
        for row in first_month[:4]
        if isinstance(row, dict)
    )

    hypothesis_steps = [
        "Собрать 10 релевантных вакансий и выделить повторяющиеся требования.",
        "Адаптировать CV под основной маршрут и отправить 5 откликов.",
        "Провести 3 тестовых интервью-скрипта и зафиксировать обратную связь.",
    ]

    resume_good = _list_items(resume_analysis.get("what_is_good"))[:6]
    resume_missing = _list_items(resume_analysis.get("what_is_missing"))[:6]
    resume_conflicts = _list_items(resume_analysis.get("inconsistencies"))[:6]
    resume_professions = _list_items(resume_analysis.get("professions"))[:6]
    resume_periods = _list_items(resume_analysis.get("periods"))[:6]
    resume_tasks = _list_items(resume_analysis.get("tasks"))[:6]
    resume_education = _list_items(resume_analysis.get("education"))[:6]
    resume_languages = _list_items(resume_analysis.get("languages"))[:6]
    resume_certificates = _list_items(resume_analysis.get("certificates"))[:6]
    resume_questions = _list_items(resume_analysis.get("clarifying_questions"))[:6]
    has_resume_module = bool(resume_analysis)
    story_echo = _build_story_echo(report)
    personal_insights = report.get("personal_insights") if isinstance(report.get("personal_insights"), list) else []
    strategy_scenarios = report.get("development_scenarios") if isinstance(report.get("development_scenarios"), list) else []
    income_forecasts = report.get("income_forecasts") if isinstance(report.get("income_forecasts"), list) else []
    horizon_plan = report.get("career_action_plan") if isinstance(report.get("career_action_plan"), dict) else {}
    insights_html = "".join(
        f"<div class='card'><p>{escape(_safe_text(item.get('insight')))}</p><p><b>Практическое следствие:</b> {escape(_safe_text(item.get('practical_consequence')))}</p></div>"
        for item in personal_insights[:5] if isinstance(item, dict)
    ) or "<p class='muted'>Инсайты требуют минимум двух подтверждённых фактов каждый.</p>"
    strategy_scenarios_html = "".join(
        f"<div class='card'><h3>{escape({'safe':'Безопасный','main':'Основной','ambitious':'Амбициозный'}.get(str(item.get('kind')), 'Сценарий'))}: {escape(_safe_text(item.get('goal')))}</h3>"
        f"<p><b>Горизонт:</b> {escape(_safe_text(item.get('horizon')))}; <b>модель:</b> {escape(_safe_text(item.get('employment_model')))}</p>"
        f"<p><b>Доход:</b> {escape(_safe_text(item.get('income_forecast')))}</p><p><b>Контрольные точки:</b> {escape(', '.join(_list_items(item.get('checkpoints'))))}</p>"
        f"<p><b>Успех:</b> {escape(_safe_text(item.get('success_criterion')))}; <b>остановка:</b> {escape(_safe_text(item.get('stop_criterion')))}</p>"
        f"<p><b>Запасной вариант:</b> {escape(_safe_text(item.get('fallback')))}</p></div>"
        for item in strategy_scenarios[:3] if isinstance(item, dict)
    )
    income_html = "".join(
        f"<tr><td>{escape(_safe_text(item.get('route_id')))}</td><td>{escape(_safe_text(item.get('country')))}</td><td>{escape(_safe_text(item.get('currency')))}</td>"
        f"<td>{escape(_safe_text(item.get('amount_type')))} / {escape(_safe_text(item.get('period')))}</td><td>{escape(_safe_text(item.get('contract_type')))}</td>"
        f"<td>{escape(_safe_text((item.get('estimates') or {}).get('conservative') if isinstance(item.get('estimates'), dict) else None))} / {escape(_safe_text((item.get('estimates') or {}).get('base') if isinstance(item.get('estimates'), dict) else None))} / {escape(_safe_text((item.get('estimates') or {}).get('optimistic') if isinstance(item.get('estimates'), dict) else None))}</td>"
        f"<td>{escape(_safe_text(item.get('data_date')))}, {escape(_safe_text(item.get('confidence')))}<br/>{escape(', '.join(_safe_text(source.get('source_name')) for source in (item.get('sources') or []) if isinstance(source, dict)))}</td></tr>"
        for item in income_forecasts if isinstance(item, dict)
    )
    def _horizon_action_html(key: str, label: str) -> str:
        action = horizon_plan.get(key)
        if not isinstance(action, dict):
            return f'<div class="card"><h3>{escape(label)}</h3><p>{escape(_safe_text(action))}</p></div>'
        return (
            f'<div class="card"><h3>{escape(label)}</h3>'
            f'<p><b>Что:</b> {escape(_safe_text(action.get("what")))}</p>'
            f'<p><b>Где и кому:</b> {escape(_safe_text(action.get("where")))} — {escape(_safe_text(action.get("audience")))}</p>'
            f'<p><b>Объём и время:</b> {escape(_safe_text(action.get("volume")))}; {escape(_safe_text(action.get("duration")))}</p>'
            f'<p><b>Успех:</b> {escape(_safe_text(action.get("success_criterion")))}</p>'
            f'<p><b>Когда менять маршрут:</b> {escape(_safe_text(action.get("change_criterion")))}</p></div>'
        )

    unicode_font = _resolve_unicode_font_path()
    font_face_css = ""
    body_font_stack = "'Arial', 'Helvetica', 'DejaVu Sans', sans-serif"
    if unicode_font:
        # Use an absolute file path to keep compatibility with both Chromium and xhtml2pdf.
        font_src = str(unicode_font.resolve()).replace("\\", "/")
        font_face_css = (
            "@font-face {"
            "font-family: 'CareerUnicode';"
            f"src: url('{font_src}') format('truetype');"
            "font-weight: normal;"
            "font-style: normal;"
            "}"
        )
        body_font_stack = "'CareerUnicode', 'Arial', 'Helvetica', 'DejaVu Sans', sans-serif"

    html = f"""
<!DOCTYPE html>
<html lang='ru'>
<head>
  <meta charset='UTF-8' />
  <style>
        {font_face_css}
    :root {{
      --paper: #ffffff;
      --ink: #1f2937;
      --muted: #4b5563;
      --accent: #0f766e;
      --line: #d1d5db;
    }}
    @page {{ size: A4; margin: 16mm; }}
        body {{ font-family: {body_font_stack}; color: var(--ink); font-size: 12px; line-height: 1.5; }}
    h1 {{ font-size: 28px; margin: 0 0 4px 0; color: #0f172a; letter-spacing: 0.2px; }}
    h2 {{ font-size: 18px; margin: 0 0 8px 0; color: var(--accent); border-left: 4px solid #8fd9cb; padding-left: 10px; }}
    h3 {{ font-size: 14px; margin: 0 0 6px 0; }}
    .page {{ page-break-after: always; }}
    .last {{ page-break-after: auto; }}
    .hero {{ border: 1px solid #cbe8e0; border-radius: 14px; background: linear-gradient(135deg, #f5fffc 0%, #eef7ff 100%); padding: 14px 16px; margin-bottom: 10px; }}
    .subtitle {{ margin: 0; color: var(--muted); font-size: 13px; }}
    .meta {{ margin-top: 10px; display: grid; grid-template-columns: 1fr 1fr; gap: 4px 10px; border-top: 1px dashed #b8d9d1; padding-top: 8px; }}
    .meta p {{ margin: 2px 0; }}
    .card {{ border: 1px solid var(--line); border-radius: 10px; padding: 10px 12px; margin-bottom: 10px; background: var(--paper); }}
        .grid2 {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px; }}
    ul {{ margin: 6px 0 0 16px; padding: 0; }}
    li {{ margin: 2px 0; }}
    .muted {{ color: var(--muted); }}
        .brand {{ display: inline-block; font-size: 11px; font-weight: bold; color: var(--accent); border: 1px solid #99f6e4; background: #f0fdfa; border-radius: 999px; padding: 3px 10px; margin-bottom: 8px; }}
        .closing-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 8px; }}
        .closing-card {{ border: 1px solid #bae6fd; border-radius: 10px; padding: 12px; background: #f8fafc; }}
        .closing-title {{ font-size: 15px; color: #0f172a; margin-bottom: 6px; }}
        .swot-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 8px; }}
        .system-note {{ border-left: 4px solid #0f766e; padding: 10px 12px; background: #f0fdfa; border-radius: 8px; margin-top: 10px; }}
        .final-conclusion {{ background: linear-gradient(135deg, #f0fdfa 0%, #eff6ff 100%); border: 2px solid #0f766e; border-radius: 14px; padding: 20px 22px; margin-top: 16px; page-break-inside: avoid; }}
        .final-conclusion-title {{ font-size: 17px; font-weight: bold; color: #0f172a; margin-bottom: 10px; border-bottom: 1px solid #99f6e4; padding-bottom: 6px; }}
        .final-conclusion-text {{ font-size: 13px; line-height: 1.7; color: #1f2937; white-space: pre-line; }}
        .action-box {{ border: 1px solid #bbf7d0; background: #f0fdf4; border-radius: 10px; padding: 14px 16px; margin-top: 12px; }}
        .action-box-title {{ font-size: 13px; font-weight: bold; color: #166534; margin-bottom: 6px; }}
        .honest-note {{ border-left: 4px solid #f59e0b; padding: 10px 14px; background: #fffbeb; border-radius: 8px; margin-top: 12px; font-size: 12px; color: #78350f; }}
        .table-scroll {{ overflow-x: auto; -webkit-overflow-scrolling: touch; }}
        table {{ width: 100%; border-collapse: collapse; min-width: 720px; }}
        th, td {{ border: 1px solid var(--line); padding: 6px; text-align: left; vertical-align: top; }}
        @media (max-width: 640px) {{ .grid2, .meta, .swot-grid, .closing-grid {{ grid-template-columns: 1fr; }} body {{ font-size: 14px; }} .page {{ page-break-after: auto; }} }}
  </style>
</head>
<body>
  <section class='page'>
                <div class='hero'>
                    <div class='brand'>NextYou</div>
                    <h1>NextYou Report</h1>
                    <p class='subtitle'>Персональная карта карьерного перехода</p>
                    <div class='meta'>
                        <p><b>Имя пользователя:</b> {escape(meta.user_name)}</p>
                        <p><b>Страна:</b> {escape(country_name)}</p>
                        <p><b>Город:</b> {escape(city)}</p>
                        <p><b>Дата:</b> {escape(meta.created_at)}</p>
                        <p><b>Режим:</b> {escape(meta.mode)}</p>
                        <p><b>Версия профиля:</b> {escape(meta.profile_version)}</p>
                    </div>
                </div>
  </section>

    <section class='page'>
        <h2>Подробный анализ по 15 блокам</h2>
        <div class='card'><h3>1. Как мы поняли вашу ситуацию</h3><p>{escape(story_echo)}</p></div>
        <div class='card'><h3>2. Ваше профессиональное ядро</h3><p>{escape(_professional_core_summary(report))}</p></div>
        <div class='card'><h3>3. Подтверждённые функции</h3><ul>{function_lines}</ul></div>
        <div class='card'><h3>4. Уровень по каждой функции</h3><ul>{level_lines}</ul></div>
        <div class='card'><h3>5. Что нельзя утверждать на основании имеющихся данных</h3><ul>{unknowns_html}</ul></div>
        <div class='card'><h3>6. Что вы хотите сохранить</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in _list_items(not_reset)[:6])}</ul></div>
        <div class='card'><h3>7. От чего вы хотите уйти</h3><ul>{refusals_html}</ul></div>
        <div class='card'><h3>8-11. Маршруты с обязательным блоком доказательств</h3>{route_evidence_html}</div>
        <div class='card'><h3>12. Требования и ограничения</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in weaknesses_items) or '<li>Данных недостаточно.</li>'}</ul></div>
        <div class='card'><h3>13. Неопределённости</h3><ul>{unknowns_html}</ul></div>
        <div class='card'><h3>14. Первый шаг до 15 минут</h3><p>{escape(_safe_text(today.get('action')))}</p></div>
        <div class='card'><h3>15. План на неделю и месяц</h3><ul>{''.join(f"<li>День {escape(str(item.get('day', '-')))}: {escape(_safe_text(item.get('task')))}</li>" for item in weekly[:7] if isinstance(item, dict)) or '<li>Данных недостаточно.</li>'}</ul></div>
        <div class='final-conclusion'><div class='final-conclusion-title'>Персональный следующий шаг</div><div class='final-conclusion-text'>{escape(_guidance_text(report))}</div></div>
    </section>

    <section class='page'>
        <h2>1. Что я услышал</h2>
        <div class='card'><h3>Кто вы сейчас</h3><p>{escape(_safe_text(digital_human.get('current_state')))}</p></div>
        <div class='card'><h3>Ключевые факты из истории и резюме (2-4)</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in heard_facts[:4])}</ul></div>
        <div class='card'><h3>Главный запрос, ресурс и ограничение</h3><ul><li><b>Главный запрос:</b> {escape(main_request)}</li><li><b>Ресурс:</b> {escape(_level_label(report.get('resource_level')))}</li><li><b>Ограничение:</b> {escape(primary_constraint)}</li></ul></div>

        <h2>2. Профессиональное ядро</h2>
        <div class='card'><h3>Ваш профессиональный капитал</h3><p>{escape(_professional_core_summary(report))}</p></div>
        <div class='card'><h3>Что не обнулилось после миграции</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in _list_items(not_reset)[:8])}</ul></div>
        <div class='card'><h3>Слои опыта</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in _list_items(experience_layers)[:6])}</ul></div>

        <h2>3. Сильные стороны и опоры</h2>
        <div class='card'><h3>Подтверждённые сильные стороны</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in strengths_items) or '<li>Данных недостаточно.</li>'}</ul></div>
        <div class='card'><h3>Источники энергии</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in _list_items(energy_sources)[:6])}</ul></div>
        <div class='card'><h3>Карьерные приоритеты</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in _list_items(career_priorities)[:6])}</ul></div>
        <div class='card'><h3>Подтверждённые компетенции</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in _list_items(competency_signals)[:6])}</ul></div>

        <h2>4. Ограничения и неизвестные</h2>
        <div class='card'><h3>Ограничения</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in weaknesses_items) or '<li>Данных недостаточно.</li>'}</ul></div>
        <div class='card'><h3>Что нужно уточнить</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in unknowns) or '<li>Критичных неизвестных сейчас нет.</li>'}</ul></div>
        <div class='swot-grid'>
            <div class='card'><h3>SWOT: Strengths</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in strengths_items) or '<li>Данных недостаточно.</li>'}</ul></div>
            <div class='card'><h3>SWOT: Weaknesses</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in weaknesses_items) or '<li>Данных недостаточно.</li>'}</ul></div>
            <div class='card'><h3>SWOT: Blind Spots</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in blind_spots) or '<li>Данных недостаточно.</li>'}</ul></div>
            <div class='card'><h3>SWOT: Important to Consider</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in important_consider) or '<li>Данных недостаточно.</li>'}</ul></div>
        </div>
        <div class='card'>
            <h3>Психологический и социальный анализ (без медицинских выводов)</h3>
            <p><b>Что можно делать самому:</b></p>
            <ul>{''.join(f'<li>{escape(x)}</li>' for x in self_help_points)}</ul>
            <p><b>Когда полезен специалист:</b></p>
            <ul>{''.join(f'<li>{escape(x)}</li>' for x in specialist_points)}</ul>
        </div>

        <h2>5. Устойчивость в период изменений</h2>
        <div class='card'><h3>Ресурс и рабочий темп</h3><p>{escape(resource_level).replace('\n', '<br/>')}</p></div>
        <div class='card'><h3>Главный риск</h3><p>{escape(_safe_text(digital_human.get('main_risk')))}</p></div>
        <div class='card'><h3>Главный страх</h3><p>{escape(_safe_text(digital_human.get('main_fear')))}</p></div>

        <h2>6. Интеграция в новой стране</h2>
        <div class='card'><h3>Состояние интеграции</h3><p>{escape(integration_level).replace('\n', '<br/>')}</p></div>
        <div class='card'><h3>Ваш главный актив</h3><p>{escape(_safe_text(digital_human.get('main_asset')))}</p></div>
        <div class='card'><h3>Скрытые активы</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in _list_items(digital_human.get('hidden_strengths'))[:6])}</ul></div>
                <div class='card'><h3>Люди и контакты</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in _list_items(social_integration.get('people'))[:6])}</ul></div>
                <div class='card'><h3>Сообщества и мосты</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in _list_items(social_integration.get('communities'))[:6])}</ul></div>
                <div class='card'><h3>Понимание рынка и возможности</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in _list_items(social_integration.get('opportunities'))[:6])}</ul></div>
  </section>

    <section class='page'>
        <h2>Что здесь легко не заметить</h2>{insights_html}
        <h2>Рынок выбранной страны и прогноз дохода</h2>
        <div class='table-scroll'><table><thead><tr><th>Маршрут</th><th>Страна</th><th>Валюта</th><th>Тип / период</th><th>Договор</th><th>Осторожно / база / оптимистично</th><th>Дата / уверенность</th></tr></thead><tbody>{income_html or '<tr><td colspan="7">Нет актуальных источников — локальные суммы не оценивались.</td></tr>'}</tbody></table></div>
        <h2>Три сценария развития</h2>{strategy_scenarios_html or '<p class="muted">Сценарии пока не сформированы.</p>'}
        <h2>План на 48 часов, 14 дней и 90 дней</h2>
        <div class='grid2'>{''.join(_horizon_action_html(key, label) for key, label in [('48_hours','Первые 48 часов'),('14_days','Первые 14 дней'),('90_days','Первые 90 дней')])}</div>
    </section>

    <section class='page'>
        <h2>7. Сравнение маршрутов</h2>
          <div class='muted'>Анализ возможностей</div>
    {''.join(possibilities) if possibilities else '<p class="muted">Данных недостаточно.</p>'}
        <div class='card'><h3>Что рынок будет проверять</h3><ul>{''.join(f'<li>{escape(x)}</li>' for x in market_questions[:10]) or '<li>Данных недостаточно.</li>'}</ul></div>
                <div class='card'><h3>Рекомендованные роли</h3>{''.join(recommendations_html) if recommendations_html else '<p class="muted">Данных недостаточно.</p>'}</div>
                <div class='card'><h3>Три сценария: безопасный, основной, амбициозный</h3>{''.join(solutions_html) if solutions_html else '<p class="muted">Данных недостаточно.</p>'}</div>
        <div class='card'><h3>Перевод опыта на язык рынка</h3>{''.join(translation_html) if translation_html else '<p class="muted">Данных недостаточно.</p>'}</div>
        <div class='card'><h3>Карьерные мосты и шаги входа</h3>{''.join(bridges_html) if bridges_html else '<p class="muted">Данных недостаточно.</p>'}</div>
        <div class='card'><h3>Риски маршрутов и барьеры</h3>{''.join(barriers_html) if barriers_html else '<p class="muted">Данных недостаточно.</p>'}</div>
        <div class='card'>
            <h3>Контекст застревания</h3>
            <ul>
                <li><b>Внешние барьеры:</b> {escape(', '.join(_list_items(barrier_landscape.get('external'))[:5]))}</li>
                <li><b>Внутренние барьеры:</b> {escape(', '.join(_list_items(barrier_landscape.get('internal'))[:5]))}</li>
                <li><b>Поведенческий риск:</b> {escape(_safe_text(barrier_landscape.get('behavioral_risk')))}</li>
                <li><b>Первое противодействие:</b> {escape(_safe_text(barrier_landscape.get('first_counter_action')))}</li>
            </ul>
        </div>
    </section>

    <section class='page'>
        <h2>8. Выбранный маршрут и первый шаг</h2>
          <div class='muted'>План действий</div>
    <div class='card'><h3>Главное решение системы</h3><p>{escape(_safe_text(decision.get('recommended_main_path')))}</p></div>
        <div class='card'><h3>Запасной маршрут</h3><p>{escape(_safe_text(decision.get('backup_path')))}</p></div>
        <div class='card'><h3>Почему именно оно</h3><p>{escape(_safe_text(decision.get('why_this_path')))}</p></div>
        <div class='card'><h3>Что не делать сейчас</h3><p>{escape(_safe_text(decision.get('avoid_for_now')))}</p></div>
        {strategy_html}
        <div class='card'><h3>Как проверить гипотезу</h3><ul>{''.join(f'<li>{escape(step)}</li>' for step in hypothesis_steps)}</ul></div>
        <div class='card'><h3>Кнопки первого шага</h3><ul>{''.join(f'<li>{escape(item)}</li>' for item in _first_step_buttons(report))}</ul></div>

        <h2>9. План на 30 дней</h2>
    <div class='card'>
      <h3>Сегодня</h3>
      <p><b>Действие:</b> {escape(_safe_text(today.get('action')))}</p>
      <p><b>Время:</b> {escape(_safe_text(today.get('timebox')))}</p>
      <p><b>Результат:</b> {escape(_safe_text(today.get('result')))}</p>
    </div>
    <div class='card'><h3>Первая неделя</h3><ul>{this_week}</ul></div>
        <div class='card'><h3>Цели на месяц</h3><ul>{this_month}</ul></div>
        <div class='card'><h3>Карта 4 недель</h3><ul>{month_rows or '<li>Данных недостаточно.</li>'}</ul></div>
        <div class='card'><h3>Недельный ритм (7 дней)</h3><ul>{''.join(f"<li>День {escape(str(item.get('day', '-')))}: {escape(_safe_text(item.get('task')))}</li>" for item in weekly[:7] if isinstance(item, dict)) or '<li>Данных недостаточно.</li>'}</ul></div>
    </section>

    <section class='page'>
        <h2>10. Анализ резюме</h2>
        {
            (
                "<div class='card'><h3>Профессии и периоды</h3><ul>"
                + ''.join(f"<li>{escape(x)}</li>" for x in resume_professions + resume_periods)
                + "</ul></div>"
                + "<div class='card'><h3>Задачи и достижения</h3><ul>"
                + ''.join(f"<li>{escape(x)}</li>" for x in resume_tasks + resume_good)
                + "</ul></div>"
                + "<div class='card'><h3>Образование, языки, сертификаты</h3><ul>"
                + ''.join(f"<li>{escape(x)}</li>" for x in resume_education + resume_languages + resume_certificates)
                + "</ul></div>"
                "<div class='card'><h3>Сильные стороны CV</h3><ul>"
                + ''.join(f"<li>{escape(x)}</li>" for x in resume_good)
                + "</ul></div>"
                + "<div class='card'><h3>Что недосказано для маршрута</h3><ul>"
                + ''.join(f"<li>{escape(x)}</li>" for x in resume_missing)
                + "</ul></div>"
                + "<div class='card'><h3>Несостыковки для уточнения</h3><ul>"
                + ''.join(f"<li>{escape(x)}</li>" for x in resume_conflicts)
                + "</ul></div>"
                + "<div class='card'><h3>Вопросы для уточнения</h3><ul>"
                + ''.join(f"<li>{escape(x)}</li>" for x in resume_questions)
                + "</ul></div>"
            )
            if has_resume_module
            else "<div class='card'><h3>Резюме не загружено</h3><p>Загрузите CV для отдельного анализа и адаптации под выбранный маршрут.</p><p><b>Кнопка:</b> «Загрузить резюме для анализа»</p></div>"
        }
    </section>

        <section class='page last'>
        <h2>11. Что может быть не так в моём выводе</h2>

        <div class='system-note'>
            <b>На чём основан этот вывод:</b> история, ответы на вопросы и данные резюме (если загружено).
            Карта меняется при новых фактах о языке, документах, приоритетах, контактах или рынке.
            Если что-то не совпадает — исправьте через кнопки ниже.
        </div>

        <div class='closing-grid' style='margin-top:12px;'>
            <div class='closing-card'>
                <div class='closing-title'>✅ Что не обнулилось — ваша опора</div>
                <ul>{''.join(f'<li>{escape(x)}</li>' for x in strengths_for_closing) or '<li>Есть устойчивые сильные стороны. Уточните в беседе со специалистом.</li>'}</ul>
            </div>
            <div class='closing-card'>
                <div class='closing-title'>⚠️ Что может изменить карту</div>
                <ul>{''.join(f'<li>{escape(x)}</li>' for x in unknowns) or '<li>Критичных неизвестных сейчас нет.</li>'}</ul>
            </div>
        </div>

        <div class='action-box' style='margin-top:14px;'>
            <div class='action-box-title'>🎯 Первый шаг — сделайте сегодня (до 15 минут)</div>
            <p style='margin:4px 0;'>{escape(first_step_today)}</p>
            <div class='action-box-title' style='margin-top:8px;'>📅 На эту неделю</div>
            <p style='margin:4px 0;'>{escape(next_step_week)}</p>
        </div>

        <div class='honest-note'>
            <b>Честно о рекомендациях специалиста:</b>
            {''.join(f' {escape(x)}' for x in specialist_hints[:2]) or ' Если маршрут понятен, но нет движения — разбор со специалистом ускоряет выход.'}
            Это не медицинская рекомендация — только карьерная и адаптационная опора.
        </div>

        <div class='closing-grid' style='margin-top:12px;'>
            <div class='closing-card'>
                <div class='closing-title'>Что можно делать самому</div>
                <ul>{''.join(f'<li>{escape(x)}</li>' for x in self_help_points[:3]) or '<li>Сделайте один короткий шаг 5-20 минут и зафиксируйте результат.</li>'}</ul>
            </div>
            <div class='closing-card'>
                <div class='closing-title'>О чем подумать со специалистом</div>
                <ul>{''.join(f'<li>{escape(x)}</li>' for x in specialist_points[:3]) or '<li>Если шаги не запускаются, полезно разобрать барьер и переформулировать маршрут.</li>'}</ul>
            </div>
        </div>

        <div class='card' style='margin-top:14px;'>
            <h3>🔎 Проверьте карту</h3>
            <ul>
                <li>Всё похоже на правду</li>
                <li>Исправить факт</li>
                <li>Изменить приоритет</li>
                <li>Не согласен с маршрутом</li>
            </ul>
        </div>

        <div class='final-conclusion' style='margin-top:18px;'>
            <div class='final-conclusion-title'>💬 Заключение</div>
            <div class='final-conclusion-text'>{escape(closing_message)}</div>
        </div>

  </section>
</body>
</html>
"""
    return re.sub(r"\n\s+", "\n", html).strip()


def _html_to_plain_text(html: str) -> str:
    # Keep a readable fallback when rich HTML->PDF engines are unavailable.
    text = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
    text = re.sub(r"</p>", "\n\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</h[1-6]>", "\n\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</li>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<li>", "- ", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = unescape(text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _render_plain_text_pdf(text: str, output_path: Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas
    except Exception as exc:
        raise RuntimeError("Fallback PDF writer requires reportlab") from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4

    font_name = "Helvetica"
    font_path = _resolve_unicode_font_path()
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("CareerUnicode", str(font_path)))
            font_name = "CareerUnicode"
        except Exception:
            font_name = "Helvetica"

    top_margin = 48
    bottom_margin = 48
    line_height = 14
    text_width_chars = 105

    y = height - top_margin
    c.setFont(font_name, 11)

    for raw_line in text.splitlines() or [""]:
        line = raw_line.rstrip()
        wrapped = textwrap.wrap(line, width=text_width_chars) if line else [""]
        for chunk in wrapped:
            if y <= bottom_margin:
                c.showPage()
                c.setFont(font_name, 11)
                y = height - top_margin
            c.drawString(40, y, chunk)
            y -= line_height

    c.save()


def html_to_pdf(html: str, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    engine = (settings.report_pdf_engine or "auto").strip().lower()
    try_playwright = engine in {"auto", "playwright"}
    try_xhtml = engine in {"auto", "xhtml2pdf"}
    errors: list[str] = []

    # Primary engine: Playwright gives the best layout fidelity for modern HTML/CSS reports.
    if try_playwright:
        try:
            from playwright.sync_api import sync_playwright

            temp_html_path = None
            try:
                with tempfile.NamedTemporaryFile(
                    mode="w",
                    suffix=".html",
                    delete=False,
                    encoding="utf-8",
                    dir=str(output_path.parent),
                ) as temp_file:
                    temp_file.write(html)
                    temp_html_path = Path(temp_file.name)

                with sync_playwright() as p:
                    browser = p.chromium.launch(headless=True)
                    page = browser.new_page(viewport={"width": 1280, "height": 1800})
                    page.set_default_timeout(45000)
                    page.goto(temp_html_path.as_uri(), wait_until="domcontentloaded")
                    page.wait_for_load_state("networkidle")
                    page.wait_for_function(
                        "() => !document.fonts || document.fonts.status === 'loaded'",
                        timeout=12000,
                    )
                    page.emulate_media(media="screen")
                    page.pdf(
                        path=str(output_path),
                        format="A4",
                        print_background=True,
                        margin={"top": "18mm", "right": "18mm", "bottom": "18mm", "left": "18mm"},
                    )
                    browser.close()
                return
            finally:
                if temp_html_path and temp_html_path.exists():
                    temp_html_path.unlink(missing_ok=True)
        except Exception as exc:
            errors.append(f"playwright: {exc}")
            if engine == "playwright":
                raise RuntimeError(f"Playwright PDF engine failed. Details: {exc}")
            # Fallback engine below keeps report generation available even without Playwright browsers.
            pass

    # Fallback engine: xhtml2pdf.
    if not try_xhtml:
        raise RuntimeError(f"Unknown REPORT_PDF_ENGINE: {engine}")

    try:
        from xhtml2pdf import pisa
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception as exc:
        errors.append(f"xhtml2pdf-import: {exc}")
        if engine == "xhtml2pdf":
            raise RuntimeError(f"xhtml2pdf engine unavailable: {exc}") from exc
        plain_text = _html_to_plain_text(html)
        _render_plain_text_pdf(plain_text, output_path)
        return

    font_path = _resolve_unicode_font_path()
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("CareerUnicode", str(font_path)))
        except Exception:
            # Keep fallback fonts if font registration fails.
            pass

    def _link_callback(uri: str, rel: str | None = None) -> str:
        if not uri:
            return ""
        parsed = urlparse(uri)
        if parsed.scheme == "file":
            resolved = unquote(parsed.path)
            if resolved.startswith("/") and len(resolved) > 2 and resolved[2] == ":":
                resolved = resolved[1:]
            return resolved
        if parsed.scheme in {"http", "https", "data"}:
            return uri
        if rel:
            return str((Path(rel).parent / uri).resolve())
        return str(Path(uri).resolve())

    # xhtml2pdf is sensitive to custom @font-face declarations in some Windows envs.
    html_for_xhtml = re.sub(r"@font-face\s*\{[^}]*\}", "", html, flags=re.IGNORECASE | re.DOTALL)
    html_for_xhtml = html_for_xhtml.replace("'CareerUnicode', ", "")

    with output_path.open("wb") as fh:
        status = pisa.CreatePDF(src=html_for_xhtml, dest=fh, encoding="utf-8", link_callback=_link_callback)
    if status.err:
        errors.append("xhtml2pdf: Failed to convert HTML to PDF")
        if engine == "xhtml2pdf":
            joined = " | ".join(errors) if errors else "unknown"
            raise RuntimeError(f"Failed to convert HTML to PDF: {joined}")
        plain_text = _html_to_plain_text(html)
        _render_plain_text_pdf(plain_text, output_path)


def generate_pdf_from_html_file_with_error(html_path: Path) -> tuple[Path | None, str]:
    try:
        html = html_path.read_text(encoding="utf-8")
    except Exception as exc:
        return None, f"html_read_error: {exc}"
    pdf_path = html_path.with_suffix(".pdf")
    try:
        html_to_pdf(html, pdf_path)
    except Exception as exc:
        return None, f"pdf_render_error: {exc}"
    return pdf_path, ""


def generate_plain_pdf_from_html_file_with_error(html_path: Path) -> tuple[Path | None, str]:
    """Force plain-text PDF generation path for diagnostics.

    This bypasses HTML engines and renders readable text PDF via reportlab.
    """
    try:
        html = html_path.read_text(encoding="utf-8")
    except Exception as exc:
        return None, f"html_read_error: {exc}"
    pdf_path = html_path.with_name(f"{html_path.stem}.plain.pdf")
    try:
        plain_text = _html_to_plain_text(html)
        _render_plain_text_pdf(plain_text, pdf_path)
    except Exception as exc:
        return None, f"plain_pdf_render_error: {exc}"
    return pdf_path, ""


def generate_pdf_report(report: dict, output_dir: str, user_name: str = "") -> Path:
    pdf_path, _, _ = generate_report_files(report, output_dir=output_dir, user_name=user_name)
    if pdf_path is None:
        raise RuntimeError("PDF generation failed")
    return pdf_path


def generate_html_report_file(report: dict, output_dir: str, user_name: str = "", profile_version: str = "") -> Path:
    meta = build_meta(report, user_name=user_name, profile_version=profile_version)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", meta.user_name)[:40] or "user"
    base_dir = Path(output_dir)
    if str(profile_version or "").strip():
        safe_version = re.sub(r"[^a-zA-Z0-9_-]+", "_", str(profile_version).strip())[:64] or "version"
        html_path = base_dir / f"career_report_{safe_name}_v_{safe_version}.html"
    else:
        html_path = base_dir / f"career_report_{safe_name}_{ts}.html"
    html = render_report_html(report, meta)
    base_dir.mkdir(parents=True, exist_ok=True)
    html_path.write_text(html, encoding="utf-8")
    return html_path


def generate_assessment_html_file(assessment, output_dir: str) -> Path:
    from services.career_assessment import render_assessment_html  # noqa: PLC0415

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    html_path = output_path / f"career_assessment_{assessment.assessment_id}.html"
    html_path.write_text(render_assessment_html(assessment), encoding="utf-8")
    return html_path


def generate_docx_report_file(
    report: dict,
    output_dir: str,
    user_name: str = "",
    profile_version: str = "",
) -> tuple[Path | None, str]:
    """Generate a DOCX version of the report. Returns (path, error_string)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as exc:
        return None, f"python-docx unavailable: {exc}"

    meta = build_meta(report, user_name=user_name, profile_version=profile_version)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", meta.user_name)[:40] or "user"
    base_dir = Path(output_dir)
    base_dir.mkdir(parents=True, exist_ok=True)
    if str(profile_version or "").strip():
        safe_version = re.sub(r"[^a-zA-Z0-9_-]+", "_", str(profile_version).strip())[:64] or "version"
        docx_path = base_dir / f"career_report_{safe_name}_v_{safe_version}.docx"
    else:
        docx_path = base_dir / f"career_report_{safe_name}_{ts}.docx"

    digital_human = report.get("digital_human", {}) if isinstance(report.get("digital_human"), dict) else {}
    decision = report.get("career_decision", {}) if isinstance(report.get("career_decision"), dict) else {}
    action_plan = report.get("action_plan", {}) if isinstance(report.get("action_plan"), dict) else {}
    today = action_plan.get("today", {}) if isinstance(action_plan.get("today"), dict) else {}
    market = report.get("market_analysis", []) if isinstance(report.get("market_analysis"), list) else []
    recommendations = report.get("career_recommendations", []) if isinstance(report.get("career_recommendations"), list) else []
    real_solutions = report.get("real_solutions", []) if isinstance(report.get("real_solutions"), list) else []
    translation = report.get("career_translation", []) if isinstance(report.get("career_translation"), list) else []
    bridges = report.get("career_bridges", []) if isinstance(report.get("career_bridges"), list) else []
    not_reset = report.get("what_not_reset", []) if isinstance(report.get("what_not_reset"), list) else []
    experience_layers = report.get("experience_layers", []) if isinstance(report.get("experience_layers"), list) else []
    career_barriers = report.get("career_barriers", []) if isinstance(report.get("career_barriers"), list) else []
    barrier_landscape = report.get("barrier_landscape", {}) if isinstance(report.get("barrier_landscape"), dict) else {}
    weekly = report.get("weekly_plan", []) if isinstance(report.get("weekly_plan"), list) else []
    social_integration = report.get("social_integration", {}) if isinstance(report.get("social_integration"), dict) else {}
    energy_sources = report.get("energy_sources", []) if isinstance(report.get("energy_sources"), list) else []
    career_priorities = report.get("career_priorities", []) if isinstance(report.get("career_priorities"), list) else []
    competency_signals = report.get("competency_signals", []) if isinstance(report.get("competency_signals"), list) else []
    resume_analysis = report.get("resume_analysis", {}) if isinstance(report.get("resume_analysis"), dict) else {}
    development = report.get("development_map", {}) if isinstance(report.get("development_map"), dict) else {}
    first_month = development.get("first_month", []) if isinstance(development.get("first_month"), list) else []
    barriers_obj = digital_human.get("barriers") if isinstance(digital_human.get("barriers"), dict) else {}
    strengths_items = list(dict.fromkeys(_list_items(not_reset)[:4] + _list_items(digital_human.get("hidden_strengths"))[:4]))[:6]
    weaknesses_items = list(dict.fromkeys(_list_items(barriers_obj.get("internal"))[:3] + _list_items(barriers_obj.get("external"))[:3]))[:6]
    story_facts = _heard_facts(report, limit=4)
    main_request = _main_request(report)
    resource_level_human = _resource_human_message(report.get("resource_level"))
    integration_level_human = _integration_human_message(report.get("integration_level"))
    primary_constraint = next((item for item in weaknesses_items if item != "-"), "данных недостаточно")
    route_scenario_labels = ["Пессимистичный", "Базовый", "Оптимистичный"]
    closing_message = _safe_text(
        report.get("closing_message"),
        "Это работа, а не испытание. Один проверяемый шаг — и карта начнёт двигаться.",
    )
    self_help_pts, specialist_pts = _psych_social_recommendation(report)
    unknowns = _unknowns_list(report)

    def _teal(r: int = 15, g: int = 118, b: int = 110) -> RGBColor:
        return RGBColor(r, g, b)

    def _add_heading(doc: "Document", text: str, level: int = 1) -> None:
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = _teal()

    def _add_bullet(doc: "Document", text: str, bold_prefix: str = "") -> None:
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            run = p.add_run(bold_prefix)
            run.bold = True
        p.add_run(_safe_text(text))

    def _add_kv(doc: "Document", key: str, value: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(f"{key}: ")
        run.bold = True
        p.add_run(_safe_text(value))

    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # Cover
        title = doc.add_heading("NextYou Report", 0)
        title.runs[0].font.color.rgb = _teal()
        doc.add_paragraph(f"Имя: {meta.user_name}  |  Страна: {meta.country}  |  Дата: {meta.created_at}  |  Режим: {meta.mode}")
        doc.add_paragraph("")

        # 1. Что я услышал
        _add_heading(doc, "1. Что я услышал", 1)
        _add_kv(doc, "Кто вы сейчас", _safe_text(digital_human.get("current_state")))
        _add_heading(doc, "Ключевые факты из истории и резюме (2-4)", 2)
        for item in story_facts:
            _add_bullet(doc, item)
        _add_kv(doc, "Главный запрос", main_request)
        _add_kv(doc, "Ресурс", _level_label(report.get("resource_level")))
        _add_kv(doc, "Ограничение", primary_constraint)
        doc.add_paragraph("")

        # 2. Профессиональное ядро
        _add_heading(doc, "2. Профессиональное ядро", 1)
        doc.add_paragraph(_professional_core_summary(report))
        _add_heading(doc, "Что не обнулилось после миграции", 2)
        for item in _list_items(not_reset)[:8]:
            _add_bullet(doc, item)
        _add_heading(doc, "Слои опыта", 2)
        for item in _list_items(experience_layers)[:6]:
            _add_bullet(doc, item)
        doc.add_paragraph("")

        # 3. Сильные стороны и опоры
        _add_heading(doc, "3. Сильные стороны и опоры", 1)
        _add_heading(doc, "Подтвержденные сильные стороны", 2)
        for item in strengths_items or ["Данных недостаточно"]:
            _add_bullet(doc, item)
        _add_heading(doc, "Источники энергии", 2)
        for item in _list_items(energy_sources)[:6]:
            _add_bullet(doc, item)
        _add_heading(doc, "Карьерные приоритеты", 2)
        for item in _list_items(career_priorities)[:6]:
            _add_bullet(doc, item)
        _add_heading(doc, "Подтвержденные компетенции", 2)
        for item in _list_items(competency_signals)[:6]:
            _add_bullet(doc, item)
        doc.add_paragraph("")

        # 4. Ограничения и неизвестные
        _add_heading(doc, "4. Ограничения и неизвестные", 1)
        _add_heading(doc, "Ограничения", 2)
        for item in weaknesses_items or ["Данных недостаточно"]:
            _add_bullet(doc, item)
        _add_heading(doc, "Что нужно уточнить", 2)
        for item in unknowns or ["Критичных неизвестных сейчас нет"]:
            _add_bullet(doc, item)
        _add_heading(doc, "SWOT: Strengths", 2)
        for item in strengths_items or ["Данных недостаточно"]:
            _add_bullet(doc, item)
        _add_heading(doc, "SWOT: Weaknesses", 2)
        for item in weaknesses_items or ["Данных недостаточно"]:
            _add_bullet(doc, item)
        _add_heading(doc, "Психологический и социальный анализ (без медицинских выводов)", 2)
        _add_heading(doc, "Что можно делать самому", 3)
        for item in self_help_pts:
            _add_bullet(doc, item)
        _add_heading(doc, "Когда полезен специалист", 3)
        for item in specialist_pts:
            _add_bullet(doc, item)
        doc.add_paragraph("")

        # 5. Устойчивость в период изменений
        _add_heading(doc, "5. Устойчивость в период изменений", 1)
        _add_kv(doc, "Ресурс и рабочий темп", resource_level_human)
        _add_kv(doc, "Главный риск", _safe_text(digital_human.get("main_risk")))
        _add_kv(doc, "Главный страх", _safe_text(digital_human.get("main_fear")))
        doc.add_paragraph("")

        # 6. Интеграция в новой стране
        _add_heading(doc, "6. Интеграция в новой стране", 1)
        _add_kv(doc, "Состояние интеграции", integration_level_human)
        _add_kv(doc, "Главный актив", _safe_text(digital_human.get("main_asset")))
        _add_heading(doc, "Скрытые активы", 2)
        for item in _list_items(digital_human.get("hidden_strengths"))[:6]:
            _add_bullet(doc, item)
        _add_heading(doc, "Люди и контакты", 2)
        for item in _list_items(social_integration.get("people"))[:6]:
            _add_bullet(doc, item)
        _add_heading(doc, "Сообщества и мосты", 2)
        for item in _list_items(social_integration.get("communities"))[:6]:
            _add_bullet(doc, item)
        _add_heading(doc, "Понимание рынка и возможности", 2)
        for item in _list_items(social_integration.get("opportunities"))[:6]:
            _add_bullet(doc, item)
        doc.add_paragraph("")

        # 7. Сравнение маршрутов
        _add_heading(doc, "7. Сравнение маршрутов", 1)
        for idx, item in enumerate(market[:3]):
            if not isinstance(item, dict):
                continue
            label = route_scenario_labels[idx] if idx < len(route_scenario_labels) else f"Вариант {idx+1}"
            _add_heading(doc, f"{label}: {_safe_text(item.get('profession'))}", 2)
            _add_kv(doc, "Соответствие", f"{item.get('fit_percent', '-')}%")
            _add_kv(doc, "Доход", _safe_text(item.get("salary_range")))
            _add_kv(doc, "Скорость входа", _safe_text(item.get("entry_speed")))
            _add_kv(doc, "Требования", ", ".join(_list_items(item.get("requirements"))[:5]))
            _add_kv(doc, "Риск", _safe_text(item.get("competition")))
        _add_heading(doc, "Рекомендованные роли", 2)
        for item in recommendations[:4]:
            if not isinstance(item, dict):
                continue
            _add_bullet(doc, f"{_safe_text(item.get('title'))}: {_safe_text(item.get('why_fit'))}")
        _add_heading(doc, "Три сценария: пессимистичный, базовый, оптимистичный", 2)
        for idx, item in enumerate(real_solutions[:3]):
            if not isinstance(item, dict):
                continue
            label = route_scenario_labels[idx] if idx < len(route_scenario_labels) else f"Вариант {idx+1}"
            _add_bullet(doc, f"{label}: {_safe_text(item.get('title'))}; {_safe_text(item.get('why'))}")
        _add_heading(doc, "Перевод опыта на язык рынка", 2)
        for item in translation[:6]:
            if not isinstance(item, dict):
                continue
            _add_bullet(doc, f"{_safe_text(item.get('source_experience'))} -> {_safe_text(item.get('market_term'))}")
        _add_heading(doc, "Карьерные мосты", 2)
        for item in bridges[:5]:
            if not isinstance(item, dict):
                continue
            _add_bullet(doc, f"{_safe_text(item.get('role'))}: {_safe_text(item.get('first_market_test'))}")
        _add_heading(doc, "Риски маршрутов и барьеры", 2)
        for item in career_barriers[:5]:
            if not isinstance(item, dict):
                continue
            _add_bullet(doc, f"{_safe_text(item.get('barrier'))}: {_safe_text(item.get('mechanism'))}")
        _add_heading(doc, "Контекст застревания", 2)
        _add_kv(doc, "Внешние барьеры", ", ".join(_list_items(barrier_landscape.get("external"))[:5]))
        _add_kv(doc, "Внутренние барьеры", ", ".join(_list_items(barrier_landscape.get("internal"))[:5]))
        _add_kv(doc, "Поведенческий риск", _safe_text(barrier_landscape.get("behavioral_risk")))
        _add_kv(doc, "Первое противодействие", _safe_text(barrier_landscape.get("first_counter_action")))
        doc.add_paragraph("")

        # 8. Выбранный маршрут и первый шаг
        _add_heading(doc, "8. Выбранный маршрут и первый шаг", 1)
        _add_kv(doc, "Маршрут", _safe_text(decision.get("recommended_main_path")))
        _add_kv(doc, "Почему", _safe_text(decision.get("why_this_path")))
        _add_kv(doc, "Запасной", _safe_text(decision.get("backup_path")))
        _add_kv(doc, "Что не делать сейчас", _safe_text(decision.get("avoid_for_now")))
        _add_kv(doc, "Первый шаг", _safe_text(today.get("action")))
        _add_kv(doc, "Время", _safe_text(today.get("timebox")))
        _add_kv(doc, "Результат", _safe_text(today.get("result")))
        _add_heading(doc, "Кнопки первого шага", 2)
        for item in _first_step_buttons(report):
            _add_bullet(doc, item)
        doc.add_paragraph("")

        # 9. План на 30 дней
        _add_heading(doc, "9. План на 30 дней", 1)
        _add_heading(doc, "Первая неделя", 2)
        for task in _list_items(action_plan.get("this_week"))[:7]:
            _add_bullet(doc, task)
        _add_heading(doc, "Цели на месяц", 2)
        for task in _list_items(action_plan.get("this_month"))[:6]:
            _add_bullet(doc, task)
        for week in first_month[:4]:
            if not isinstance(week, dict):
                continue
            _add_heading(doc, f"Неделя {week.get('week', '-')}: {_safe_text(week.get('focus'))}", 2)
            for task in _list_items(week.get("tasks", []))[:4]:
                _add_bullet(doc, task)
            doc.add_paragraph(f"Результат недели: {_safe_text(week.get('output'))}")
        _add_heading(doc, "Недельный ритм (7 дней)", 2)
        if weekly:
            for item in weekly[:7]:
                if isinstance(item, dict):
                    _add_bullet(doc, f"День {item.get('day', '-')}: {_safe_text(item.get('task'))}")
        else:
            _add_bullet(doc, "Данных недостаточно")
        doc.add_paragraph("")

        # 10. Анализ резюме
        _add_heading(doc, "10. Анализ резюме", 1)
        if resume_analysis:
            _add_heading(doc, "Профессии и периоды", 2)
            for item in _list_items(resume_analysis.get("professions"))[:6] + _list_items(resume_analysis.get("periods"))[:6]:
                _add_bullet(doc, item)
            _add_heading(doc, "Задачи и достижения", 2)
            for item in _list_items(resume_analysis.get("tasks"))[:6] + _list_items(resume_analysis.get("what_is_good"))[:6]:
                _add_bullet(doc, item)
            _add_heading(doc, "Образование, языки, сертификаты", 2)
            for item in _list_items(resume_analysis.get("education"))[:6] + _list_items(resume_analysis.get("languages"))[:6] + _list_items(resume_analysis.get("certificates"))[:6]:
                _add_bullet(doc, item)
            _add_heading(doc, "Что недосказано для маршрута", 2)
            for item in _list_items(resume_analysis.get("what_is_missing"))[:6]:
                _add_bullet(doc, item)
            _add_heading(doc, "Несостыковки для уточнения", 2)
            for item in _list_items(resume_analysis.get("inconsistencies"))[:6]:
                _add_bullet(doc, item)
            _add_heading(doc, "Вопросы для уточнения", 2)
            for item in _list_items(resume_analysis.get("clarifying_questions"))[:6]:
                _add_bullet(doc, item)
        else:
            doc.add_paragraph("Резюме не загружено. Загрузите CV для отдельного анализа под маршрут.")
            _add_bullet(doc, "Кнопка: Загрузить резюме для анализа")
        doc.add_paragraph("")

        # 11. Что может быть не так
        _add_heading(doc, "11. Что может быть не так в моем выводе", 1)
        doc.add_paragraph(
            "Карта меняется при новых данных о языке, документах, резюме, приоритетах или рынке."
        )
        if unknowns:
            _add_heading(doc, "Что нужно уточнить:", 2)
            for u in unknowns:
                _add_bullet(doc, u)
        _add_heading(doc, "Что делать самому:", 2)
        for pt in self_help_pts:
            _add_bullet(doc, pt)
        _add_heading(doc, "О чем подумать со специалистом:", 2)
        for pt in specialist_pts:
            _add_bullet(doc, pt)
        _add_heading(doc, "Кнопки валидации карты", 2)
        for item in ["Все похоже на правду", "Исправить факт", "Изменить приоритет", "Не согласен с маршрутом"]:
            _add_bullet(doc, item)
        doc.add_paragraph("")

        # Closing
        closing_p = doc.add_paragraph()
        closing_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = closing_p.add_run("💬 Заключение")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = _teal()
        doc.add_paragraph(closing_message)

        doc.save(str(docx_path))
        return docx_path, ""
    except Exception as exc:
        return None, f"docx_render_error: {exc}"


def generate_pdf_from_html_file(html_path: Path) -> Path | None:
    pdf_path, _error = generate_pdf_from_html_file_with_error(html_path)
    return pdf_path


def generate_report_files(
    report: dict,
    output_dir: str,
    user_name: str = "",
    profile_version: str = "",
) -> tuple[Path | None, Path, Path | None]:
    html_path = generate_html_report_file(
        report,
        output_dir=output_dir,
        user_name=user_name,
        profile_version=profile_version,
    )
    pdf_path = generate_pdf_from_html_file(html_path)
    docx_path, _ = generate_docx_report_file(
        report,
        output_dir=output_dir,
        user_name=user_name,
        profile_version=profile_version,
    )
    return pdf_path, html_path, docx_path


def generate_report_payload(user_id: str, report: dict, base_url: str, output_dir: str, user_name: str = "") -> dict[str, str]:
    pdf_path = generate_pdf_report(report, output_dir=output_dir, user_name=user_name)
    filename = pdf_path.name
    base = base_url.rstrip("/")
    return {
        "telegram_summary": build_telegram_summary(report),
        "pdf_url": f"{base}/{filename}",
    }
